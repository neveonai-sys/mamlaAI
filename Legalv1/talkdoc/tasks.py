import base64
import csv
import io
import os
import logging
import zipfile
from datetime import datetime
from xml.etree import ElementTree as ET
from bson import ObjectId
from celery import shared_task
from core.init_clients import get_mongo_client
from core.llm_client import vision_complete
from .search import ensure_index
from .chunk import split_into_chunks

EMBED_MODEL = os.getenv("RAG_EMBED_MODEL", "text-embedding-3-large")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.webp', '.bmp', '.gif', '.tif', '.tiff'}
CSV_EXTENSIONS = {'.csv'}
XLSX_EXTENSIONS = {'.xlsx'}

def _mongo():
    return get_mongo_client()['legaldb']


def _extension(filename: str) -> str:
    return os.path.splitext((filename or '').lower())[1]


def _format_table(rows, title=None):
    clean_rows = []
    for row in rows or []:
        normalized = [str(cell or '').replace('\n', ' ').strip() for cell in row]
        if any(normalized):
            clean_rows.append(normalized)

    if not clean_rows:
        return ''

    width = max(len(row) for row in clean_rows)
    padded_rows = [row + [''] * (width - len(row)) for row in clean_rows]
    header = padded_rows[0]
    body = padded_rows[1:]

    lines = []
    if title:
        lines.append(title)
    lines.append('| ' + ' | '.join(header) + ' |')
    lines.append('| ' + ' | '.join(['---'] * width) + ' |')
    for row in body:
        lines.append('| ' + ' | '.join(row) + ' |')
    return '\n'.join(lines)


def _extract_docx_text(raw_bytes: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(raw_bytes))
    parts = []

    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    if paragraphs:
        parts.append('\n'.join(paragraphs))

    for index, table in enumerate(document.tables, start=1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        table_text = _format_table(rows, title=f'[Table {index}]')
        if table_text:
            parts.append(table_text)

    return '\n\n'.join(parts)


def _extract_pdf_text(raw_bytes: bytes) -> str:
    import pdfplumber
    import pypdfium2 as pdfium

    parts = []
    pdfium_doc = pdfium.PdfDocument(raw_bytes)
    with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            page_text = (page.extract_text() or '').strip()
            if page_text:
                parts.append(f'[Page {page_number}]\n{page_text}')

            found_table = False
            for table_index, table in enumerate(page.extract_tables() or [], start=1):
                table_text = _format_table(table, title=f'[Page {page_number} Table {table_index}]')
                if table_text:
                    found_table = True
                    parts.append(table_text)

            # Scanned PDFs often have no text layer. Rasterize those pages and use the
            # same multimodal OCR path used for standalone image uploads.
            if not page_text and not found_table:
                ocr_text = _extract_pdf_page_image_text(pdfium_doc, page_number - 1)
                if ocr_text:
                    parts.append(f'[Page {page_number} OCR]\n{ocr_text}')

    pdfium_doc.close()

    return '\n\n'.join(parts)


def _extract_csv_text(raw_bytes: bytes) -> str:
    decoded = raw_bytes.decode('utf-8', errors='ignore')
    rows = list(csv.reader(io.StringIO(decoded)))
    return _format_table(rows, title='[CSV Table]')


def _extract_xlsx_text(raw_bytes: bytes) -> str:
        namespace = {'main': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}

        def column_index(cell_ref: str) -> int:
            letters = ''.join(char for char in cell_ref if char.isalpha())
            idx = 0
            for char in letters:
                idx = idx * 26 + (ord(char.upper()) - 64)
            return max(idx - 1, 0)

        def load_shared_strings(archive: zipfile.ZipFile):
            if 'xl/sharedStrings.xml' not in archive.namelist():
                return []
            root = ET.fromstring(archive.read('xl/sharedStrings.xml'))
            values = []
            for node in root.findall('main:si', namespace):
                text_parts = [part.text or '' for part in node.findall('.//main:t', namespace)]
                values.append(''.join(text_parts).strip())
            return values

        def resolve_cell_value(cell, shared_strings):
            cell_type = cell.attrib.get('t')
            value_node = cell.find('main:v', namespace)
            inline_node = cell.find('main:is', namespace)
            if inline_node is not None:
                return ''.join(part.text or '' for part in inline_node.findall('.//main:t', namespace)).strip()
            if value_node is None or value_node.text is None:
                return ''
            raw = value_node.text.strip()
            if cell_type == 's':
                try:
                    return shared_strings[int(raw)]
                except Exception:
                    return raw
            return raw

        with zipfile.ZipFile(io.BytesIO(raw_bytes)) as archive:
            workbook_root = ET.fromstring(archive.read('xl/workbook.xml'))
            rel_root = ET.fromstring(archive.read('xl/_rels/workbook.xml.rels'))
            shared_strings = load_shared_strings(archive)

            rel_map = {}
            for relation in rel_root.findall('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                rel_map[relation.attrib.get('Id')] = relation.attrib.get('Target')

            parts = []
            sheets = workbook_root.findall('main:sheets/main:sheet', namespace)
            for sheet in sheets:
                sheet_name = sheet.attrib.get('name', 'Sheet')
                relationship_id = sheet.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                target = rel_map.get(relationship_id)
                if not target:
                    continue
                target_path = f"xl/{target}" if not target.startswith('xl/') else target
                sheet_root = ET.fromstring(archive.read(target_path))
                rows = []
                for row_node in sheet_root.findall('.//main:sheetData/main:row', namespace):
                    row_values = []
                    for cell in row_node.findall('main:c', namespace):
                        idx = column_index(cell.attrib.get('r', 'A1'))
                        while len(row_values) <= idx:
                            row_values.append('')
                        row_values[idx] = resolve_cell_value(cell, shared_strings)
                    if any(value.strip() for value in row_values):
                        rows.append(row_values)
                table_text = _format_table(rows, title=f'[{sheet_name}]')
                if table_text:
                    parts.append(table_text)

            return '\n\n'.join(parts)


def _image_to_data_url(raw_bytes: bytes, mimetype: str) -> str:
    encoded = base64.b64encode(raw_bytes).decode('utf-8')
    return f'data:{mimetype};base64,{encoded}'


def _extract_image_text(raw_bytes: bytes, filename: str, mimetype: str) -> str:
    prompt = (
        'You are performing OCR and layout extraction for a legal document image. '
        'Extract all readable text in natural reading order. '
        'If the image contains a table, reconstruct it as a markdown table. '
        'Preserve headings, numbering, party names, dates, monetary values, stamps, signatures, and short handwritten notes when legible. '
        'Return only the extracted document text with no summary or commentary. '
        'If a span is unreadable, use [illegible].'
    )
    return vision_complete(
        prompt=prompt,
        image_data_url=_image_to_data_url(raw_bytes, mimetype),
        app_scenario='talkdoc:rag',
        temperature=0.1,
        max_tokens=3000,
    )


def _extract_pdf_page_image_text(pdfium_doc, page_index: int) -> str:
    page = pdfium_doc.get_page(page_index)
    try:
        bitmap = page.render(scale=2.0)
        try:
            image = bitmap.to_pil()
            buffer = io.BytesIO()
            image.save(buffer, format='PNG')
            return _extract_image_text(buffer.getvalue(), f'page-{page_index + 1}.png', 'image/png')
        finally:
            bitmap.close()
    finally:
        page.close()


def _extract_text(raw_bytes: bytes, filename: str, mimetype: str) -> str:
    extension = _extension(filename)

    if extension == '.docx':
        return _extract_docx_text(raw_bytes)
    if extension == '.pdf':
        return _extract_pdf_text(raw_bytes)
    if extension == '.txt':
        return raw_bytes.decode('utf-8', errors='ignore')
    if extension in CSV_EXTENSIONS or mimetype in {'text/csv', 'application/csv'}:
        return _extract_csv_text(raw_bytes)
    if extension in XLSX_EXTENSIONS or mimetype in {'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'}:
        return _extract_xlsx_text(raw_bytes)
    if extension in IMAGE_EXTENSIONS or mimetype.startswith('image/'):
        return _extract_image_text(raw_bytes, filename, mimetype)

    raise ValueError(f'Unsupported file type for TalkDoc extraction: {extension or mimetype or filename}')


def _friendly_ingest_error(stage, error):
    message = str(error).strip().splitlines()[0]
    if stage == 'gridfs':
        return 'The stored file could not be retrieved for processing.'
    if stage == 'extract':
        if 'Unsupported file type' in message:
            return 'This file type is not supported for Talk To Docs extraction yet.'
        if 'No extractable text found in document' in message:
            return 'No readable text was found in this document.'
        return 'Text extraction failed for this document.'
    if stage == 'chunk':
        return 'The extracted text could not be prepared for indexing.'
    if stage == 'embed':
        return 'Embedding failed while preparing this document.'
    if stage == 'index':
        return 'Search indexing failed for this document.'
    return 'Document processing failed. Please try uploading it again.'


def _mark_failed(db, doc, stage, error, logger):
    detail = str(error)
    logger.error(f"[INGEST][ERROR] stage={stage} detail={detail}")
    db['rag_documents'].update_one(
        {"_id": doc["_id"]},
        {"$set": {
            "status": "failed",
            "ingest_stage": "failed",
            "error": _friendly_ingest_error(stage, error),
            "error_detail": detail,
            "updated_at": datetime.utcnow(),
        }},
    )

def embed_texts(texts):
    """
    Lightweight OpenAI embeddings call. Replace with your existing client if preferred.
    """
    from openai import OpenAI
    
    client = OpenAI(api_key=OPENAI_API_KEY)
    resp = client.embeddings.create(model=EMBED_MODEL, input=texts)
    return [d.embedding for d in resp.data]

@shared_task(bind=True, max_retries=3)
def ingest_document(self, doc_id: str):
    logger = logging.getLogger('django')
    db = _mongo()
    doc = db['rag_documents'].find_one({"_id": ObjectId(doc_id)})
    if not doc:
        logger.error(f"[INGEST] Document {doc_id} not found in DB.")
        return

    import traceback
    logger.info(f"[INGEST] Starting ingestion for doc_id: {doc_id}")
    try:
        db['rag_documents'].update_one(
            {"_id": doc["_id"]},
            {"$set": {"status": "processing", "ingest_stage": "extracting", "error": "", "error_detail": "", "updated_at": datetime.utcnow()}},
        )

        # 1) download from GridFS
        file_id = doc["storage"].get("file_id")
        if not file_id:
            logger.error(f"[INGEST][ERROR] No file_id found in document storage for doc_id: {doc_id}")
            raise Exception("No file_id found in document storage")
        logger.info(f"[INGEST] GridFS file_id: {file_id} type: {type(file_id)}")
        # Ensure file_id is ObjectId
        if not isinstance(file_id, ObjectId):
            try:
                file_id = ObjectId(file_id)
            except Exception as id_err:
                logger.error(f"[INGEST][ERROR] Could not convert file_id to ObjectId: {id_err}")
                _mark_failed(db, doc, 'gridfs', id_err, logger)
                raise Exception(f"Could not convert file_id to ObjectId: {id_err}")
        from gridfs import GridFS
        gridfs_api = GridFS(get_mongo_client()["legaldb"], collection="talkdoc_files")
        try:
            file_obj = gridfs_api.get(ObjectId(file_id))
            logger.info(f"[INGEST] Successfully fetched file from GridFS for doc_id: {doc_id}")
        except Exception as gridfs_err:
            logger.error(f"[INGEST][ERROR] File with id {file_id} not found in GridFS: {gridfs_err}\n{traceback.format_exc()}")
            _mark_failed(db, doc, 'gridfs', gridfs_err, logger)
            raise Exception(f"File with id {file_id} not found in GridFS: {gridfs_err}")
        raw_bytes = file_obj.read()

        # 2) extract text directly
        text = ""
        filename = doc.get("name_original") or doc["name_stored"]
        mimetype = doc.get('mimetype') or 'application/octet-stream'
        try:
            logger.info(f"[INGEST] Extracting text from file: {filename}")
            text = _extract_text(raw_bytes, filename, mimetype)
            logger.info(f"[INGEST] Text extraction complete. Length: {len(text)} chars")
            if not text.strip():
                raise ValueError("No extractable text found in document")
        except Exception as extract_err:
            logger.error(f"[INGEST][ERROR] Text extraction failed: {extract_err}\n{traceback.format_exc()}")
            _mark_failed(db, doc, 'extract', extract_err, logger)
            raise

        # 3) chunk
        try:
            db['rag_documents'].update_one(
                {"_id": doc["_id"]},
                {"$set": {"status": "processing", "ingest_stage": "chunking", "updated_at": datetime.utcnow()}},
            )
            logger.info(f"[INGEST] Chunking text for doc_id: {doc_id}")
            chunks = split_into_chunks(text, max_tokens=700, overlap=80)
            logger.info(f"[INGEST] Chunking complete. Num chunks: {len(chunks)}")
            if not chunks:
                raise ValueError("No chunks generated from extracted text")
        except Exception as chunk_err:
            logger.error(f"[INGEST][ERROR] Chunking failed: {chunk_err}\n{traceback.format_exc()}")
            _mark_failed(db, doc, 'chunk', chunk_err, logger)
            raise

        # 4) embed
        try:
            db['rag_documents'].update_one(
                {"_id": doc["_id"]},
                {"$set": {"status": "processing", "ingest_stage": "embedding", "updated_at": datetime.utcnow()}},
            )
            logger.info(f"[INGEST] Embedding chunks for doc_id: {doc_id}")
            vecs = embed_texts([c["text"] for c in chunks])
            logger.info(f"[INGEST] Embedding complete. Num vectors: {len(vecs)}")
        except Exception as embed_err:
            logger.error(f"[INGEST][ERROR] Embedding failed: {embed_err}\n{traceback.format_exc()}")
            _mark_failed(db, doc, 'embed', embed_err, logger)
            raise

        # 5) index to OpenSearch
        try:
            db['rag_documents'].update_one(
                {"_id": doc["_id"]},
                {"$set": {"status": "processing", "ingest_stage": "indexing", "updated_at": datetime.utcnow()}},
            )
            logger.info(f"[INGEST] Indexing chunks to OpenSearch for doc_id: {doc_id}")
            cli = ensure_index()
            actions = []
            for i, c in enumerate(chunks):
                body = {
                    "chunk_id": f"{doc_id}_{i}",
                    "user_id": doc["user_id"],
                    "doc_id": str(doc["_id"]),
                    "matter": doc.get("matter", {}),
                    "name_stored": doc["name_stored"],
                    "page": None,
                    "text": c["text"],
                    "vector": vecs[i],
                    "created_at": datetime.utcnow()
                }
                actions.append({"index": {"_index": os.getenv("RAG_OS_INDEX","rag_chunks_v1")}})
                actions.append(body)
            if actions:
                cli.bulk(body=actions, refresh=True)
            logger.info(f"[INGEST] Indexing complete for doc_id: {doc_id}")
        except Exception as index_err:
            logger.error(f"[INGEST][ERROR] Indexing failed: {index_err}\n{traceback.format_exc()}")
            _mark_failed(db, doc, 'index', index_err, logger)
            raise

        db['rag_documents'].update_one({"_id": doc["_id"]},
                                       {"$set": {"status": "indexed", "ingest_stage": "indexed", "pages": None, "updated_at": datetime.utcnow()}})
        logger.info(f"[INGEST] Ingestion complete for doc_id: {doc_id}")
    except Exception as e:
        logger.error(f"[INGEST][FATAL] Ingestion failed for doc_id: {doc_id}: {e}\n{traceback.format_exc()}")
        _mark_failed(db, doc, 'fatal', e, logger)
        raise
