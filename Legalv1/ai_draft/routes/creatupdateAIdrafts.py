from bson import ObjectId
from core.llm_client import chat_complete
from talkdoc.tasks import _extract_image_text
from ai_draft.citation_grounding import build_grounding_block
from ai_draft.drafting.classify import DraftContext, classify
from ai_draft.drafting.prompt_builder import (
    build_draft_system_prompt,
    build_location_string,
    target_max_tokens,
)
from ai_draft.drafting.draft_validator import (
    build_correction_message,
    parse_draft_payload,
    validate,
)
from users.routes.encryption import encrypt_field, decrypt_field

import os
import math
from io import BytesIO
from docx import Document
from pymongo import ASCENDING, DESCENDING, ReturnDocument
from core.init_clients import get_mongo_client, get_mongo_db
from pdfminer.high_level import extract_text as extract_text_from_pdf
import datetime
import json
import traceback
import logging
logger = logging.getLogger('django')


#: Output schema for drafting calls. The advisory object gives assumptions and
#: drafting notes somewhere to live — defect #4 is unsatisfiable in the legacy
#: bare array, and Phase 1 showed the model responding to that absence by
#: writing them as prose above the JSON, which lost the entire draft.
#: `parse_draft_payload` accepts both shapes, so this is safe to flip either way.
DRAFT_SCHEMA = 'advisory'

#: One correction turn, never more. Worst case is two calls per generation.
MAX_CORRECTION_TURNS = 1

#: Rank a DraftResult so a correction turn that came back worse is discarded.
#: Ordering: usable beats fatal; fewer errors beats more; then more sections.
def _result_rank(result):
    if result is None:
        return (-1, 0, 0)
    if result.fatal:
        return (0, 0, 0)
    return (1, -len(result.errors), len(result.sections))


#: Cache key for the draft read path.
#:
#: Versioned, and owned here rather than repeated as an f-string at each of the
#: seven call sites that read or invalidate it. The `v2` payload is the whole
#: response — sections, status, history and advisories — because the `v1` shape
#: cached sections alone and reconstructed the rest with hard-coded values,
#: which made a still-generating draft read back as completed.
def draft_cache_key(session_id) -> str:
    return f'draft_payload:v2:{session_id}'


#: Free-text keys inside an advisory record. Everything else in the record
#: (severity, confirm_with_client) is an enum or a bool and carries no client
#: facts, so it stays queryable in plaintext.
_ADVISORY_TEXT_KEYS = ('assumption', 'why', 'issue', 'recommendation')


def _encrypt_advisories(items):
    """
    Fernet-wrap the free text of assumptions / drafting notes.

    A parallel of `_encrypt_sections` rather than a reuse of it: advisories have
    a different shape, and running them through the section helper would make
    them look like sections to anything that later iterates the collection.
    """
    out = []
    for item in items or []:
        if not isinstance(item, dict):
            continue
        rec = dict(item)
        for key in _ADVISORY_TEXT_KEYS:
            if isinstance(rec.get(key), str) and rec[key]:
                rec[key] = encrypt_field(rec[key])
        out.append(rec)
    return out


def _decrypt_advisories(items):
    """Inverse of `_encrypt_advisories` — safe on already-plaintext data."""
    out = []
    for item in items or []:
        if not isinstance(item, dict):
            continue
        rec = dict(item)
        for key in _ADVISORY_TEXT_KEYS:
            if isinstance(rec.get(key), str) and rec[key]:
                rec[key] = decrypt_field(rec[key])
        out.append(rec)
    return out


def _encrypt_sections(sections):
    """
    Encrypt the 'content' field of each draft section before it goes into
    Mongo (Privacy Policy Section 7 / Sensitive Personal Data: "Case details,
    legal documents"). Also encrypts each section's 'history' snapshots,
    which store prior 'content' values. Returns a new list — never mutates
    the input in place, since callers often still need the plaintext
    version for an immediate API response.
    """
    result = []
    for s in sections or []:
        s = dict(s)
        if 'content' in s:
            s['content'] = encrypt_field(s['content'])
        if s.get('history'):
            s['history'] = [
                {**h, 'content': encrypt_field(h['content'])} if 'content' in h else h
                for h in s['history']
            ]
        result.append(s)
    return result


def _decrypt_sections(sections):
    """Inverse of _encrypt_sections — safe to call on already-plaintext data."""
    result = []
    for s in sections or []:
        s = dict(s)
        if 'content' in s:
            s['content'] = decrypt_field(s['content'])
        if s.get('history'):
            s['history'] = [
                {**h, 'content': decrypt_field(h['content'])} if 'content' in h else h
                for h in s['history']
            ]
        result.append(s)
    return result

class CreateupdatefetchAIdrafts:
    def __init__(self,user_id):
        self.user_id = user_id

    def get_mongo_client_db(self):
        mongo = get_mongo_client()
        if not mongo:
            return ''
        db = get_mongo_db()
        collection = db['aidrafts_complete_data']
        # Indexes are created via scripts/optimize_database_indexes.py
        # No need to create them on every request
        return collection

    def get_total_drafts_count(self):
        try:
            total_drafts = self.get_mongo_client_db().count_documents({'user_id': self.user_id})
            return total_drafts
        except Exception as e:
            logger.error(f" error at ---- get_total_drafts_count ---> {traceback.format_exc()}")
            return 0

    def start_new_session(self, user_query, draft_for, location={}, language='English',
                          *, document_type=None):
        """Create session and generate draft synchronously (legacy method - use async version)"""
        try:
            logger.info(f"[start_session] Received user_query: {user_query}")
            time_now = datetime.datetime.now(datetime.timezone.utc)

            # Classify once, here, and persist it. Every later operation on this
            # session — refine, section edit, validation — reads the stored
            # context rather than re-deriving it, so a refine cannot end up on a
            # different playbook than the draft it is revising.
            #
            # `draft_for` is NOT passed as a type hint: it is case/client
            # association, and `normalize_type_hint` refuses that shape anyway.
            ctx = classify(user_query, document_type)
            logger.info(
                "[start_session] classified as %s (branch=%s, via %s)",
                ctx.doc_type, ctx.branch, ctx.source,
            )

            # Create a new session
            session = {
                'user_id': self.user_id,
                'user_query': user_query,
                'draft_for': draft_for,
                'ai_suggested_update_count': 0,
                'location': location,
                'language': language,
                'draft_context': ctx.to_dict(),
                'draft_sections': [],
                'conversation_history': [],
                'created_on': time_now,
                'last_updated_on': time_now,
                'status': 'generating'
            }
            session_id = self.get_mongo_client_db().insert_one(session).inserted_id
            self.generate_draft(session_id, user_query, location, language, ctx=ctx)
            logger.info(f"[start_session] New session created with session_id: {session_id}")
            return session_id
        except Exception as err:
            logger.error(f" error at ---- start_new_session ---> {traceback.format_exc()}")
            return ''

    def start_new_session_without_ai(self, user_query, draft_for, location={}, language='English',
                                     *, document_type=None):
        """Create session WITHOUT generating draft (for async processing)"""
        try:
            logger.info(f"[start_session_async] Received user_query: {user_query}")
            time_now = datetime.datetime.now(datetime.timezone.utc)

            # Classify here too, so the Celery worker inherits the same context
            # rather than re-classifying and possibly disagreeing with the
            # session the user is already looking at.
            ctx = classify(user_query, document_type)

            session = {
                'user_id': self.user_id,
                'user_query': user_query,
                'draft_for': draft_for,
                'ai_suggested_update_count': 0,
                'location': location,
                'language': language,
                'draft_context': ctx.to_dict(),
                'draft_sections': [],
                'conversation_history': [],
                'created_on': time_now,
                'last_updated_on': time_now,
                'status': 'generating'  # Will be updated by async task
            }
            session_id = self.get_mongo_client_db().insert_one(session).inserted_id
            logger.info(f"[start_session_async] New session created: {session_id}. AI generation will be async.")
            return session_id
        except Exception as err:
            logger.error(f"Error at start_new_session_without_ai: {traceback.format_exc()}")
            return ''

    def auto_save_initial_draft(self, session_id, draft_name, draft_sections):
        """
        Pushes a first “Untitled …” saved_draft inside the session.
        Returns draft metadata.
        """
        now = datetime.datetime.now(datetime.timezone.utc)
        draft_id = str(ObjectId())
        self.get_mongo_client_db().update_one(
            {'_id': ObjectId(session_id)},
            {'$push': {'saved_drafts': {
                'draft_id'  : draft_id,
                'draft_name': draft_name,
                'sections'  : _encrypt_sections(draft_sections),
                'saved_at'  : now
            }},
             '$set':  {'last_updated_on': now}}
        )
        return {
            'draft_id': draft_id,
            'saved_at': now,
            'last_updated_on': now,
        }

    def backfill_initial_saved_draft(self, session_id):
        """
        Copy the session's generated sections into its first saved_draft.

        Async generation creates the saved_draft row up front, while the
        sections are still empty, so the draft appears in the user's list the
        moment they ask for it. Once the worker finishes, that snapshot has to
        catch up — otherwise opening the draft from the sidebar (which reads the
        saved snapshot, not the live session) shows an empty document even
        though generation succeeded.

        Only the FIRST saved_draft is touched, and only while it is still empty:
        a user who has since saved their own revision must never have it
        overwritten by the generator.
        """
        try:
            session = self.get_mongo_client_db().find_one(
                {'_id': ObjectId(session_id)},
                {'saved_drafts': 1, 'draft_sections': 1},
            )
            if not session:
                return False

            saved = session.get('saved_drafts') or []
            if not saved or saved[0].get('sections'):
                return False        # nothing to backfill, or already populated

            sections = session.get('draft_sections') or []
            if not sections:
                return False

            now = datetime.datetime.now(datetime.timezone.utc)
            # `draft_sections` is already encrypted at rest and saved_drafts
            # uses the same shape, so this is a copy, not a re-encryption.
            self.get_mongo_client_db().update_one(
                {'_id': ObjectId(session_id),
                 'saved_drafts.draft_id': saved[0].get('draft_id')},
                {'$set': {'saved_drafts.$.sections': sections,
                          'saved_drafts.$.saved_at': now,
                          'last_updated_on': now}},
            )
            return True
        except Exception:
            # A failed backfill must not fail the draft: the session itself has
            # the sections, and the workspace reads those.
            logger.error(f"[backfill_initial_saved_draft] {traceback.format_exc()}")
            return False


    # def update_location_for_draft_creation(self, session_id, state, district):
    #     """
    #         -- initially was used, now shifted the code flow, not needed anymore
    #     """
    #     try:
    #         # Update session with location
    #         time_now = datetime.datetime.now(datetime.timezone.utc)
    #         self.get_mongo_client_db().update_one(
    #             {'_id': ObjectId(session_id)},
    #             {'$set': {'location': {'state': state, 'district': district, 'last_updated_on': time_now}}}
    #         )

    #         logger.info(f"[set_location] Updated session {session_id} with location: {state}, {district}")

    #         # Retrieve updated session
    #         session = self.get_mongo_client_db().find_one({'_id': ObjectId(session_id)})
    #         if not session:
    #             logger.error(f"[set_location] Session {session_id} not found.")
    #             return {'mssg': False}

    #         # filter applied that if loading template sections are already present, so no need to generate again...
    #         if not len(session.get("draft_sections")):
    #             # Generate draft
    #             self.generate_draft(session_id, session['user_query'], state, district)

    #         logger.info(f"[set_location] Draft generated for session {session_id}")

    #         return {'mssg': True}
    #     except Exception as err:
    #         logger.error(f" error at ---- update_location_for_draft_creation ---> {traceback.format_exc()}")
    #         return {'mssg': False}

    def fetch_existing_template_text(self, draft_file_name, draft_type):
        """
        Retrieve the document from the 'draft_content_data' collection
        that matches the given 'draft_file_name' and 'draft_type',
        and return the 'content' field if found.
        """
        query = {
            "filename": draft_file_name,  # The 'filename' you store in Mongo
            "draft_type": draft_type
        }

        mongo = get_mongo_client()
        if not mongo:
            return ''
        db = get_mongo_db()

        document = db['draft_content_data'].find_one(query)
        logger.info(f"fetch_existing_template_text ----> query: {query}, document: {document}")

        if document:
            return document.get("content", "")
        else:
            return None


    def generate_draft(self, session_id, user_query, location, language, *, ctx=None):
        logger.info(f"[generate_draft] Generating draft for session_id: {session_id} ----location >>> {location}")

        # Classify inline when the caller did not, so no entry point can bypass
        # the playbook by forgetting to pass a context.
        if ctx is None:
            ctx = classify(user_query)
            logger.info(
                "[generate_draft] no context supplied; classified as %s (branch=%s)",
                ctx.doc_type, ctx.branch,
            )

        playbook = ctx.playbook
        system_prompt = build_draft_system_prompt(
            ctx,
            language=language,
            location=location,
            # Notices have no precedent anywhere in draftdocs/, so their worked
            # example is hand-authored on the playbook. Retrieval from the
            # corpus lands in Phase 4 and will feed this same argument.
            exemplar=playbook.inline_exemplar,
            schema=DRAFT_SCHEMA,
        )
        max_tokens = target_max_tokens(ctx)

        # The prompt is static per document type; logging it on every generation
        # buried the actual user facts in noise and leaked them at INFO.
        logger.info(
            "[generate_draft] doc_type=%s branch=%s max_tokens=%d prompt_chars=%d",
            ctx.doc_type, ctx.branch, max_tokens, len(system_prompt),
        )

        messages = [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': f'Draft the document. The instructions are:\n\n{user_query}'},
        ]

        result = self._generate_and_validate(
            messages, ctx, user_query=user_query, max_tokens=max_tokens,
            session_id=session_id,
        )

        if result is None or result.fatal:
            # Previously this wrote an empty `draft_sections` and never touched
            # `status`, so the workspace polled a session that would never
            # complete. Recording the failure is what kills the perpetual
            # spinner — it does not need Celery.
            self._write_draft_result(session_id, None, ctx, status='failed')
            return

        self._write_draft_result(session_id, result, ctx, status='completed')

    def _generate_and_validate(self, messages, ctx, *, user_query, max_tokens,
                               session_id=None, scenario='ai_draft:generate'):
        """
        Call the model, parse with the repair ladder, validate, and spend at most
        one correction turn on the errors found.

        Returns the best `DraftResult` obtained, or None if every attempt failed
        to produce a usable document. The correction turn is only worth taking
        for deterministic errors — a statute the branch forbids, a missing
        mandatory section, an express instruction dropped without a word.
        Warnings are logged and shipped.
        """
        conversation = list(messages)
        best = None

        for attempt in range(MAX_CORRECTION_TURNS + 1):
            try:
                raw, meta = chat_complete(
                    messages=conversation,
                    app_scenario=scenario,
                    temperature=0.3,  # Low temp → reliable JSON structure
                    max_tokens=max_tokens,
                    return_usage=True,
                )
            except Exception:
                logger.error("[generate_draft] LLM call failed on attempt %d: %s",
                             attempt + 1, traceback.format_exc())
                return best

            finish_reason = (meta or {}).get('finish_reason')
            result = parse_draft_payload(raw)
            if not result.fatal:
                validate(result, ctx, user_query=user_query, finish_reason=finish_reason)

            logger.info(
                "[generate_draft] session=%s attempt=%d doc_type=%s finish_reason=%s %s",
                session_id, attempt + 1, ctx.doc_type, finish_reason, result.summary(),
            )

            # Keep the better of the two attempts. A correction turn can come
            # back worse than what it was correcting, and shipping the worse one
            # because it happened to be last would be a regression the user sees.
            if best is None or _result_rank(result) > _result_rank(best):
                best = result

            if result.fatal:
                if attempt >= MAX_CORRECTION_TURNS:
                    break
                conversation = conversation + [
                    {'role': 'assistant', 'content': (str(raw) or '')[:2000]},
                    {'role': 'user', 'content': (
                        'That response could not be parsed. Return ONLY the JSON described '
                        'in the output format — no prose before it, no commentary after it, '
                        'no markdown fences.'
                    )},
                ]
                continue

            if not result.errors or attempt >= MAX_CORRECTION_TURNS:
                break

            conversation = conversation + [
                {'role': 'assistant', 'content': (str(raw) or '')[:12000]},
                {'role': 'user', 'content': build_correction_message(result, ctx)},
            ]

        return best

    def _write_draft_result(self, session_id, result, ctx, *, status):
        """
        Persist sections, advisories and validation metadata in one write.

        `draft_sections` / `original_draft` keep their existing shape and
        encryption exactly. Advisories go into NEW siblings so they can never be
        mistaken for draft sections — they cannot reach the DOCX/PDF export
        paths, section reordering, or revert, all of which iterate
        `draft_sections` alone.
        """
        now = datetime.datetime.now(datetime.timezone.utc)
        update = {'status': status, 'last_updated_on': now}

        if result is None:
            update['draft_sections'] = []
            update['original_draft'] = []
            update['validation'] = {
                'checked_on': now, 'fatal': True, 'findings': [],
                'doc_type': ctx.doc_type, 'branch': ctx.branch,
            }
        else:
            self._stamp_section_ids(result.sections)
            encrypted = _encrypt_sections(result.sections)
            update['draft_sections'] = encrypted
            update['original_draft'] = encrypted
            update['draft_assumptions'] = _encrypt_advisories(result.assumptions)
            update['draft_notes'] = _encrypt_advisories(result.drafting_notes)
            update['validation'] = {
                'checked_on': now,
                'fatal': False,
                'repaired': result.repaired,
                'schema': result.schema,
                'doc_type': ctx.doc_type,
                'branch': ctx.branch,
                # Findings are quality metadata, not client content: codes and
                # severities only. The `message` field quotes offending draft
                # text, so it stays out of Mongo in plaintext.
                'findings': [
                    {'code': f.code, 'severity': f.severity, 'section_name': f.section_name}
                    for f in result.findings
                ],
            }

        self.get_mongo_client_db().update_one(
            {'_id': ObjectId(session_id)}, {'$set': update}
        )
        logger.info("[generate_draft] Session %s written with status=%s", session_id, status)


    def _stamp_section_ids(self, sections):
        """Give every section the id the storage and edit paths key on."""
        for section in sections or []:
            section['section_id'] = str(ObjectId())
        return sections

    def parse_draft_into_sections(self, draft_content):
        """
        Legacy entry point, now a wrapper over `parse_draft_payload`.

        Kept because several callers (the file-upload generators, `tasks.py`, and
        tests) want just the sections and nothing else. The repair ladder,
        schema coercion and section cleaning all live in `draft_validator` so
        there is exactly one parser in the product — the four-way prompt drift
        this codebase already suffered is not a mistake worth repeating for
        parsing.

        Advisories returned by the advisory schema are discarded here by design:
        a caller that wants them uses `parse_draft_payload` directly.
        """
        result = parse_draft_payload(draft_content)
        if result.fatal:
            logger.error("[parse_draft_into_sections] unparseable: %s",
                         '; '.join(f.message for f in result.findings))
            return []
        if result.repaired:
            logger.warning("[parse_draft_into_sections] recovered via %s (%d sections)",
                           result.repaired, len(result.sections))
        return self._stamp_section_ids(result.sections)

    def update_specific_section_of_the_draft(self, session_id, section_id, section_name, content):
        try:
            session = self.get_mongo_client_db().find_one({'_id': ObjectId(session_id)})
            section = next((s for s in session['draft_sections'] if s['section_id'] == section_id), None)
            if section:
                # Save the current content in history
                history_entry = {
                    'section_name': section['section_name'],
                    'content': section['content'],
                    'timestamp': datetime.datetime.now(datetime.timezone.utc)
                }
                self.get_mongo_client_db().update_one(
                    {'_id': ObjectId(session_id), 'draft_sections.section_id': section_id},
                    {'$push': {'draft_sections.$.history': history_entry}}
                )
                logger.info(f"update_specific_section_of_the_draft ========>>> {history_entry} || updateddd")
            else:
                return {'mssg': False}
            time_now = datetime.datetime.now(datetime.timezone.utc)
            # Update the section in the session
            result = self.get_mongo_client_db().update_one(
                {'_id': ObjectId(session_id), 'draft_sections.section_id': section_id},
                {'$set': {
                    'draft_sections.$.section_name': section_name,
                    'draft_sections.$.content': encrypt_field(content),
                    'last_updated_on': time_now
                }}
            )
            return {'mssg': True}
        except Exception as e:
            logger.error(f"[update_specific_section_of_the_draft]  ============>>>>>>: {traceback.format_exc()}")
            return {'mssg': False}

    def delete_specific_section_of_the_draft(self, session_id, section_id):
        try:
            time_now = datetime.datetime.now(datetime.timezone.utc)
            # Remove the section from the session
            result = self.get_mongo_client_db().update_one(
                {'_id': ObjectId(session_id)},
                {'$pull': {'draft_sections': {'section_id': section_id}},
                 '$set': {'last_updated_on': time_now}}
            )

            if result.modified_count == 0:
                logger.error(f"[delete_section] No matching section to delete for session_id: {session_id}, section_id: {section_id}")
                return {'mssg': False}

            logger.info(f"[delete_section] Section {section_id} deleted successfully from session {session_id}.")
            return {'mssg': True}
        except Exception as e:
            logger.error(f"[delete_specific_section_of_the_draft]  ============>>>>>>: {traceback.format_exc()}")
            return {'mssg': False}

    def update_content_using_AI_with_user_input(self, session_id, section_id, suggestion):
        try:
            session = self.get_mongo_client_db().find_one({'_id': ObjectId(session_id)})
            # all_sections is ciphertext (raw Mongo read) — decrypt before any
            # of it goes into the LLM prompt below.
            all_sections = _decrypt_sections(session.get('draft_sections', []))
            section = next(
                (s for s in all_sections if s['section_id'] == section_id), None
            )

            # Build a brief summary of the other sections for context
            other_sections_context = '\n'.join(
                f'  - {s["section_name"]}: {(s.get("content") or "")[:300].strip()}{"..." if len(s.get("content") or "") > 300 else ""}'
                for s in all_sections if s['section_id'] != section_id and (s.get('content') or '').strip()
            )
            context_block = (
                f"\nFor context, the rest of the draft contains these sections:\n{other_sections_context}\n"
                if other_sections_context else ''
            )

            # Fetch-then-inject citation grounding: if the instruction asks for a
            # citation, resolve it against the live e-SCR portal *before* the LLM
            # writes anything, and hand it verified data (or an explicit
            # "could not verify" instruction) rather than let it invent one.
            citation_block = build_grounding_block(suggestion) or ''

            system_prompt = f"""You are refining one section of a legal draft.{context_block}
The section to update is titled "{section['section_name']}":

{section['content']}

Update ONLY this section, incorporating the user's instructions. Ensure the content aligns with the rest of the draft and complies with Indian legal standards.

Return ONLY the updated section content. Do not include any headings, labels, preamble, or extra commentary.{citation_block}"""

            # Short-term working memory: last 3 exchanges of AI conversation for this section,
            # so follow-up instructions ("make it shorter") are aware of what was asked before.
            conversation_history = session.get('conversation_history', [])
            prior_turns = [
                {'role': entry.get('role', 'user'), 'content': entry.get('content', '')}
                for entry in conversation_history
                if entry.get('section_id') == section_id
            ][-6:]

            messages = [{'role': 'system', 'content': system_prompt}, *prior_turns, {'role': 'user', 'content': suggestion}]

            updated_content, _usage = chat_complete(
                messages=messages,
                app_scenario='ai_draft:update_section',
                temperature=0.4,
                max_tokens=2000,
                return_usage=True,
            )

            # Long-term audit log: persist this exchange on the session document.
            time_now = datetime.datetime.now(datetime.timezone.utc)
            self.get_mongo_client_db().update_one(
                {'_id': ObjectId(session_id)},
                {'$push': {'conversation_history': {'$each': [
                    {'role': 'user', 'content': suggestion, 'section_id': section_id, 'timestamp': time_now},
                    {'role': 'assistant', 'content': updated_content, 'section_id': section_id, 'timestamp': time_now},
                ]}}},
            )

            return {'mssg': updated_content, 'usage': _usage}
        except Exception as e:
            logger.error(f"[update_content_using_AI_with_user_input]  ============>>>>>>: {traceback.format_exc()}")
            return {'mssg': False}

    def add_new_section_in_existing_draft(self, session_id, section_name, content):
        try:
            section_id = str(ObjectId())
            plaintext_section = {
                'section_id': section_id,
                'section_name': section_name,
                'content': content
            }
            new_section = {**plaintext_section, 'content': encrypt_field(content)}
            time_now = datetime.datetime.now(datetime.timezone.utc)
            # Add the new section to the session
            result = self.get_mongo_client_db().update_one(
                {'_id': ObjectId(session_id)},
                {'$push': {'draft_sections': new_section},
                 '$set': {'last_updated_on': time_now}}
            )

            if result.modified_count == 0:
                logger.error(f"[add_section] Failed to add section to session {session_id}.")
                return {'mssg': False}

            logger.info(f"[add_section] Section {section_id} added successfully to session {session_id}.")
            return {'mssg': plaintext_section}
        except Exception as e:
            logger.error(f"[add_new_section_in_existing_draft]  ============>>>>>>: {traceback.format_exc()}")
            return {'mssg': False}

    def prepare_content_for_download(self, session_id):
        try:
            # Retrieve session
            session = self.get_mongo_client_db().find_one({'_id': ObjectId(session_id)})
            if not session:
                logger.error(f"[download_draft] Session {session_id} not found.")
                return {'mssg': False} 

            # draft_sections is ciphertext (raw Mongo read) — must decrypt
            # before compiling, otherwise the downloaded .docx would contain
            # Fernet tokens instead of the actual legal document text.
            draft_sections = _decrypt_sections(session.get('draft_sections', []))
            if not draft_sections:
                logger.error(f"[download_draft] No draft sections found for session {session_id}.")
                return {'mssg': False}

            logger.info(f"[download_draft] Compiling draft sections for session {session_id}.")

            # Compile all sections into one document
            document = Document()
            for section in draft_sections:
                section_name = section.get('section_name', 'Untitled Section')
                content = section.get('content', '')

                document.add_heading(section_name, level=1)
                document.add_paragraph(content)

                logger.info(f"[download_draft] Added section: {section_name}")

            # Save document to a BytesIO stream

            doc_io = BytesIO()
            document.save(doc_io)
            doc_io.seek(0)
            return {'mssg': doc_io}
        except Exception as e:
            logger.error(f"[prepare_content_for_download]  ============>>>>>>: {traceback.format_exc()}")
            return {'mssg': False}

    def adjust_section_position_in_draft(self, session_id,draft_sections):
        try:
            # Ensure that only the section_ids and their order are updated
            new_order = [{'section_id': s['section_id']} for s in draft_sections]

            # Fetch current sections
            session = self.get_mongo_client_db().find_one({'_id': ObjectId(session_id)})
            current_sections = session['draft_sections']

            # Create a mapping from section_id to section data
            section_map = {s['section_id']: s for s in current_sections}

            # Reorder sections
            reordered_sections = [section_map[s['section_id']] for s in new_order]

            time_now = datetime.datetime.now(datetime.timezone.utc)
            # Update the session with the new order
            self.get_mongo_client_db().update_one(
                {'_id': ObjectId(session_id)},
                {'$set': {'draft_sections': reordered_sections, 'last_updated_on': time_now}},
            )
            return {'mssg': True}
        except Exception as e:
            logger.error(f"[adjust_section_position_in_draft]  ============>>>>>>: {traceback.format_exc()}")
            return {'mssg': False}

    def retrieve_history_of_section_of_draft_if_updated(self, session_id, section_id):
        try:
            # Retrieve session and section
            session = self.get_mongo_client_db().find_one({'_id': ObjectId(session_id)})
            if not session:
                return {'mssg': False}

            section = next(
                (s for s in session['draft_sections'] if s['section_id'] == section_id), None
            )
            if not section:
                return {'mssg': False}

            history = section.get('history', [])
            history = [
                {**h, 'content': decrypt_field(h['content'])} if 'content' in h else h
                for h in history
            ]
            return {'mssg': history}
        except Exception as e:
            logger.error(f"[retrieve_history_of_section_of_draft_if_updated]  ============>>>>>>: {traceback.format_exc()}")
            return {'mssg': False}

    def retrieve_sections_of_draft(self, session_id):
        try:
            # Retrieve session
            session = self.get_mongo_client_db().find_one({'_id': ObjectId(session_id)})
            if not session:
                logger.error(f"[get_draft_sections] Session {session_id} not found.")
                return {'mssg': False, 'status': 'not_found'}

            draft_sections = _decrypt_sections(session.get('draft_sections', []))
            final_count = session.get('ai_suggested_update_count', 0)
            conversation_history = session.get('conversation_history', [])

            # A session with no sections and no recorded status is one that
            # failed before Phase 2 started writing `status` — the perpetual
            # spinner. Report it as failed so the workspace can offer a retry
            # instead of polling forever.
            status = session.get('status')
            if not status:
                status = 'completed' if draft_sections else 'failed'
            elif status == 'generating' and draft_sections:
                status = 'completed'

            logger.info(f"[get_draft_sections] Retrieved {len(draft_sections)} sections for session {session_id}. Status: {status}")
            return {
                'mssg': draft_sections,
                'ai_suggested_update_count': final_count,
                'status': status,
                'conversation_history': conversation_history,
                # Advisories travel beside the sections, never inside them.
                'assumptions': _decrypt_advisories(session.get('draft_assumptions', [])),
                'drafting_notes': _decrypt_advisories(session.get('draft_notes', [])),
            }
        except Exception as e:
            logger.error(f"[retrieve_sections_of_draft]  ============>>>>>>: {traceback.format_exc()}")
            return {'mssg': False, 'ai_suggested_update_count': 0, 'status': 'error',
                    'conversation_history': [], 'assumptions': [], 'drafting_notes': []}

    def retrieve_single_section_from_session(self, session_id, section_id):
        try:
            draft_section = self.get_mongo_client_db().find_one(
                        {
                            '_id': ObjectId(session_id),  # Match the session by session_id
                            'user_id': self.user_id,
                            'draft_sections.section_id': section_id  # Match the section_id within the draft_sections array
                        },
                        {
                            'draft_sections.$': 1  # Project the matched section from draft_sections
                        }
                    )
            latest_content = decrypt_field(draft_section['draft_sections'][0]['content'])
            logger.info(f"retrieve_single_section_from_session ========== <><><><><><><><> {latest_content}")
            return {'mssg': latest_content}
        except Exception as e:
            logger.error(f"[retrieve_single_section_from_session]  ============>>>>>>: {traceback.format_exc()}")
            return {'mssg': False}

    def revert_draft_changes_to_intial_stage(self, session_id):
        try:
            session = self.get_mongo_client_db().find_one({'_id': ObjectId(session_id)})
            if not session or 'original_draft' not in session:
                return {'mssg': False}
            time_now = datetime.datetime.now(datetime.timezone.utc)
            self.get_mongo_client_db().update_one(
                {'_id': ObjectId(session_id)},
                {'$set': {'draft_sections': session['original_draft'], 'last_updated_on': time_now}}
            )
            return {'mssg': True}
        except Exception as e:
            logger.error(f"[revert_draft_changes_to_intial_stage]  ============>>>>>>: {traceback.format_exc()}")
            return {'mssg': False}


    def save_semi_filled_drafts(self, session_id, draft_name, draft_sections, draft_id=None):
        """
        If draft_id is given, try to update that existing saved draft.
        If not found or not given, insert a new one.
        """
        try:
            session_object_id = ObjectId(session_id)

            # Validate the session
            session = self.get_mongo_client_db().find_one({'_id': session_object_id, 'user_id': self.user_id})
            if not session:
                logger.error(f"[save_semi_filled_drafts] Session not found for session_id: {session_id}")
                return {'mssg': False}

            # Make sure 'saved_drafts' is in the session
            if 'saved_drafts' not in session:
                self.get_mongo_client_db().update_one(
                    {'_id': session_object_id},
                    {'$set': {'saved_drafts': []}}
                )

            time_now = datetime.datetime.now(datetime.timezone.utc)
            # draft_sections arrives as plaintext from the caller — encrypt
            # once, reuse for whichever branch below actually runs.
            encrypted_sections = _encrypt_sections(draft_sections)

            if draft_id:
                # Overwrite the existing entry if it exists
                result = self.get_mongo_client_db().update_one(
                    {
                        '_id': session_object_id,
                        'user_id': self.user_id,
                        'saved_drafts.draft_id': draft_id
                    },
                    {
                        '$set': {
                            'saved_drafts.$.draft_name': draft_name,
                            'saved_drafts.$.sections': encrypted_sections,
                            'saved_drafts.$.saved_at': time_now,
                            'last_updated_on': time_now,
                        }
                    }
                )
                if result.modified_count > 0:
                    logger.info(f"[save_semi_filled_drafts] Updated existing saved draft: {draft_id}")
                    return {
                        'mssg': True,
                        'draft_id': draft_id,
                        'saved_at': time_now,
                        'last_updated_on': time_now,
                    }

                # If no match found, we proceed to create new 
                # or you can return an error if you'd prefer.

            # Otherwise, create a new saved draft
            new_draft_id = str(ObjectId())
            saved_draft = {
                'draft_id': new_draft_id,
                'draft_name': draft_name,
                'sections': encrypted_sections,
                'saved_at': time_now
            }
            self.get_mongo_client_db().update_one(
                {'_id': session_object_id, 'user_id': self.user_id},
                {
                    '$push': {'saved_drafts': saved_draft},
                    '$set': {'last_updated_on': time_now},
                }
            )
            logger.info(f"[save_semi_filled_drafts] Created new saved draft: {new_draft_id}")
            return {
                'mssg': True,
                'draft_id': new_draft_id,
                'saved_at': time_now,
                'last_updated_on': time_now,
            }

        except Exception as e:
            logger.error(f"[save_semi_filled_drafts] Exception: {traceback.format_exc()}")
            return {'mssg': False}


    # def get_saved_draft_list(self):
    #     try:
    #         # Fetch all sessions for the user
    #         sessions = self.get_mongo_client_db().find({'user_id': self.user_id}, {'saved_drafts': 1})

    #         saved_drafts = []
    #         for session in sessions:
    #             for draft in session.get('saved_drafts', []):
    #                 saved_drafts.append({
    #                     'draft_name': draft['draft_name'],
    #                     'draft_id': draft['draft_id'],
    #                     'created_on': draft['saved_at'],
    #                     'last_updated_on': draft.get('last_updated_on', draft['saved_at']),
    #                     'session_id': str(session['_id'])
    #                 })
    #         return {'mssg': saved_drafts}
    #     except Exception as e:
    #         logger.error(f"[get_saved_draft_list]  ============>>>>>>: {traceback.format_exc()}")
    #         return {'mssg': False}

    def get_saved_draft_list(self, page, page_size, filter_query):
        try:
            # Calculate total count for pagination
            total_count = self.get_mongo_client_db().count_documents(filter_query)
            logger.info(f"get_saved_draft_list ============== total_count ===============>.>>>>> {total_count}")
            # Calculate total pages
            total_pages = math.ceil(total_count / page_size) if page_size > 0 else 1

            # Ensure page is within bounds
            if page < 1:
                page = 1
            elif page > total_pages and total_pages > 0:
                page = total_pages

            # Define the aggregation pipeline
            pipeline = [
                        {'$match': filter_query},
                        {'$unwind': '$saved_drafts'},
                        {'$project': {
                            'draft_name': '$saved_drafts.draft_name',
                            'draft_id': {'$toString': '$saved_drafts.draft_id'},
                            'draft_for': '$draft_for',  # Extract draft_for from the top-level document
                            'created_on': {
                                '$dateToString': {
                                    'format': '%Y-%m-%dT%H:%M:%SZ',
                                    'date': '$created_on'
                                }
                            },
                            'last_updated_on': {
                                '$dateToString': {
                                    'format': '%Y-%m-%dT%H:%M:%SZ',
                                    'date': '$last_updated_on'
                                }
                            },
                            'session_id': {'$toString': '$_id'}
                        }},
                        {'$sort': {'created_on': -1}},  # Adjust sort as needed
                        {'$facet': {
                            'paginatedResults': [
                                {'$skip': (page - 1) * page_size},
                                {'$limit': page_size}
                            ],
                            'totalCount': [
                                {'$count': 'count'}
                            ]
                        }}
                    ]

            aggregation_result = list(self.get_mongo_client_db().aggregate(pipeline))

            if not aggregation_result:
                saved_drafts = []
                total_count = 0
            else:
                paginated_results = aggregation_result[0].get('paginatedResults', [])
                saved_drafts = []
                for draft in paginated_results:
                    saved_drafts.append({
                        'draft_name': draft.get('draft_name', 'N/A'),
                        'draft_id': draft.get('draft_id', 'N/A'),
                        'created_on': draft.get('created_on', datetime.datetime.now(datetime.timezone.utc).isoformat() + 'Z'),
                        'last_updated_on': draft.get('last_updated_on', datetime.datetime.now(datetime.timezone.utc).isoformat() + 'Z'),
                        'session_id': draft.get('session_id', 'N/A'),
                        'draft_for': draft.get('draft_for', {})
                    })
                #total_count = aggregation_result[0].get('totalCount', [{'count': 0}])[0].get('count', 0)
                total_count = aggregation_result[0].get('totalCount', [{'count': 0}])
                #logger.info(f"total_count ==========>>>> {total_count}")
                if len(total_count):
                    total_count = total_count[0].get('count', 0)
                else:
                    total_count = 0

            # Calculate total pages
            total_pages = math.ceil(total_count / page_size) if page_size > 0 else 1
            #logger.info(f"get_saved_draft_list ============== saved_drafts ===============>.>>>>> {saved_drafts[-1]}")

            return {
                'saved_drafts': saved_drafts,
                'pagination': {
                    'current_page': page,
                    'page_size': page_size,
                    'total_pages': total_pages,
                    'total_count': total_count
                }
            }
        except Exception as e:
            logger.error(f"[get_saved_draft_list]  ============>>>>>>: {traceback.format_exc()}")
            return {'err': False}

    def delete_saved_draft(self, session_id, draft_id):
        try:
            try:
                session_object_id = ObjectId(session_id)
                draft_object_id = ObjectId(draft_id)
            except Exception as e:
                logger.error(f"Invalid session_id or draft_id format.")
                {'mssg': False}

            # Remove the draft from saved_drafts
            result = self.get_mongo_client_db().update_one(
                {'_id': session_object_id, 'user_id':self.user_id},
                {'$pull': {'saved_drafts': {'draft_id': draft_id}}}
            )
            return {'mssg': True}
        except Exception as e:
            logger.error(f"[delete_saved_draft]  ============>>>>>>: {traceback.format_exc()}")
            return {'mssg': False}

    def load_saved_draft_details(self, session_id, draft_id):
        try:
            try:
                session_object_id = ObjectId(session_id)
            except Exception as e:
                logger.error(f"Invalid session_id format: {session_id}")
                return {'error': 'Invalid session_id format.'}

            # Fetch the session document
            session = self.get_mongo_client_db().find_one({'user_id':self.user_id, '_id': session_object_id})

            if not session:
                logger.error(f"Session not found for session_id: {session_id}")
                return {'error': 'Session not found.'}

            # Find the saved draft
            saved_draft = next((draft for draft in session.get('saved_drafts', [])
                                if draft['draft_id'] == draft_id), None)

            if not saved_draft:
                logger.error(f"Draft not found for draft_id: {draft_id}")
                return {'error': 'Draft not found.'}

            return {'draft_sections': _decrypt_sections(saved_draft['sections'])}

        except Exception as e:
            logger.error(f"Exception in load_saved_draft: {traceback.format_exc()}")
            return {'error': 'An error occurred while loading the draft.'}


    def save_draft_from_template(self, draft_type, draft_sections, draft_for=None):
        try:
            if draft_for is None:
                draft_for = {}
            time_now = datetime.datetime.now(datetime.timezone.utc)
            encrypted_sections = _encrypt_sections(draft_sections)
            # Create a new session
            session = {
                'user_id': self.user_id,
                'user_query': draft_type,
                'draft_for': draft_for,
                'location': {
                    'last_updated_on': time_now
                },
                'draft_sections': encrypted_sections,
                'conversation_history': [],
                'created_on': time_now,
                'last_updated_on': time_now,
                'original_draft': encrypted_sections
            }

            # Insert the session into MongoDB
            result = self.get_mongo_client_db().insert_one(session)
            session_id = str(result.inserted_id)
            return {"mssg":session_id}
        except Exception as e:
            logger.error(f"Exception in load_saved_draft: {traceback.format_exc()}")
            return {'error': 'An error occurred while loading the draft.'}

    # def extract_text_from_file(self,file_path):
    #     """
    #     Extract text from a file (PDF, DOC, TXT).
    #     """
    #     text = textract.process(file_path).decode('utf-8')
    #     return text

    # def generate_draft_sections_from_template(self,file_text, draft_type, state, district):
    #     """
    #     Generate draft sections from the extracted text.
    #     You can customize this function to parse the text and create sections.
    #     """
    #     # For simplicity, we'll create one section with the entire text
    #     draft_sections = [{
    #         'section_id': str(ObjectId()),
    #         'section_name': 'Template Content',
    #         'content': file_text
    #     }]
    #     return draft_sections

    def split_text_into_sections(self, text):
        """
        Split the generated text into sections.
        """
        # For simplicity, let's assume sections are separated by headings
        # Customize this function based on your needs
        lines = text.split('\n')
        sections = []
        current_section = {'section_id': str(ObjectId()), 'section_name': '', 'content': ''}
        for line in lines:
            if line.isupper():
                # Start of a new section
                if current_section['section_name']:
                    sections.append(current_section)
                    current_section = {'section_id': str(ObjectId()), 'section_name': line, 'content': ''}
                else:
                    current_section['section_name'] = line
            else:
                current_section['content'] += line + '\n'
        if current_section['section_name']:
            sections.append(current_section)
        return sections

    def insert_draft_session_for_casedocument(self, draft_sections, draft_for, language='English', user_description=None):
        try:
            # Validate 'draft_for'
            if not isinstance(draft_for, dict):
                logger.warning("draft_for is not a dictionary.")
                draft_for = {}

            # Allowed keys
            allowed_keys = {'caseid', 'clientid', 'caseid_with_clientid'}

            # Filter out unwanted keys and ensure values are lists
            draft_for_filtered = {k: v for k, v in draft_for.items() if k in allowed_keys and isinstance(v, list) and len(v) > 0}

            if 'personal' in draft_for.keys():
                draft_for_filtered['personal']='Y'
            # logger.info(f"insert_draft_session_for_casedocument ---  draft_for_filtered  ->>>>> {draft_for_filtered} ====== {draft_for}")
            # Check how many keys are populated
            # populated_keys = [k for k, v in draft_for_filtered.items() if v]
            # Create a new session
            encrypted_sections = _encrypt_sections(draft_sections)
            session = {
                'user_id': self.user_id,
                'user_query': user_description or 'case_document',
                'draft_for': draft_for_filtered,
                'language': language,
                'location': {},  # No state and district needed
                'draft_sections': encrypted_sections,
                'conversation_history': [],
                'created_on': datetime.datetime.now(datetime.timezone.utc),
                'last_updated_on': datetime.datetime.now(datetime.timezone.utc),
                'original_draft': encrypted_sections
            }

            # Insert the session into MongoDB
            result = self.get_mongo_client_db().insert_one(session)
            session_id = str(result.inserted_id)
            return {"mssg":session_id}
        except Exception as e:
            logger.error(f"[generate_draft_sections_with_gpt]  ============>>>>>>: {traceback.format_exc()}")
            return {"mssg":False}

    def extract_text_from_file(self, file_stream, file_name):
        """
        Extract text from a file-like object based on its type.
        Supports PDF, DOCX, TXT.
        """
        try:
            if file_name.endswith('.pdf'):
                # Use pdfminer.six
                text = extract_text_from_pdf(file_stream)
                return text
            elif file_name.endswith('.docx'):
                # Use python-docx
                document = Document(file_stream)
                text = '\n'.join([paragraph.text for paragraph in document.paragraphs])
                return text
            elif file_name.endswith('.txt'):
                # For text files
                text = file_stream.read().decode('utf-8')
                return text
            elif file_name.endswith(('.png', '.jpg', '.jpeg')):
                # OCR via the shared vision-LLM extractor used by TalkDoc
                mimetype = 'image/png' if file_name.endswith('.png') else 'image/jpeg'
                return _extract_image_text(file_stream.getvalue(), file_name, mimetype)
            elif file_name.endswith('.doc'):
                # For .doc files, we need to use external tools or convert to DOCX
                logger.error("Unsupported file format: .doc")
                return None
            else:
                logger.error("Unsupported file format.")
                return None
        except Exception as e:
            logger.error(f"Exception in extract_text_from_file: {traceback.format_exc()}")
            return None

    def generate_draft_sections_from_template(self, file_text, draft_type, *, ctx=None):
        """
        Generate draft sections from an existing template or an uploaded file.

        The template text is the SOURCE, not the format authority: it goes in the
        source block, while the playbook skeleton and the statute policy still
        govern. Several corpus forms predate the 2023 codes, so a template that
        cites the CrPC must not license the draft to do the same.
        """
        if ctx is None:
            # `draft_type` is a genuine type label on this path (the template
            # picker's folder name), unlike `draft_for` elsewhere.
            ctx = classify(str(draft_type or ''), draft_type)
            if ctx.is_generic and file_text:
                ctx = classify(str(file_text)[:2000], draft_type)

        system_prompt = build_draft_system_prompt(
            ctx,
            source_text=str(file_text or '')[:20000],
        )
        logger.info(
            "[generate_from_template] doc_type=%s branch=%s draft_type=%r",
            ctx.doc_type, ctx.branch, draft_type,
        )

        try:
            generated_text = chat_complete(
                messages=[
                    {'role': 'system', 'content': system_prompt},
                    {'role': 'user', 'content': (
                        f'Draft a {ctx.playbook.label} based on the source material above, '
                        'adapting its structure to the mandatory skeleton.'
                    )},
                ],
                app_scenario='ai_draft:generate_from_tpl',
                temperature=0.3,
                max_tokens=target_max_tokens(ctx),
            ).strip()

            # Parse the generated text into draft sections
            draft_sections = self.parse_draft_into_sections(generated_text)

            return draft_sections

        except Exception as e:
            logger.error(f"Exception in generate_draft_sections_from_template: {traceback.format_exc()}")
            return None

    def generate_draft_sections_with_gpt(self, file_text, language='English', user_description=None,
                                         *, ctx=None):
        """
        Generate draft sections from a case document.

        Classification prefers the user's own description of what they want, and
        only falls back to sniffing the document itself — a case file describes
        the matter, not necessarily the instrument the user is asking for.
        """
        if ctx is None:
            ctx = classify(str(user_description or ''))
            if ctx.is_generic and file_text:
                ctx = classify(str(file_text)[:2000])

        instruction = 'Draft the document from the case material above.'
        if user_description:
            instruction += f'\n\nThe user\'s instructions:\n\n{user_description}'

        system_prompt = build_draft_system_prompt(
            ctx,
            language=language,
            source_text=str(file_text or '')[:20000],
        )
        logger.info(
            "[generate_from_case] doc_type=%s branch=%s has_description=%s",
            ctx.doc_type, ctx.branch, bool(user_description),
        )

        try:
            generated_text = chat_complete(
                messages=[
                    {'role': 'system', 'content': system_prompt},
                    {'role': 'user', 'content': instruction},
                ],
                app_scenario='ai_draft:generate_from_case',
                temperature=0.3,
                max_tokens=target_max_tokens(ctx),
            ).strip()

            # Parse the generated text into draft sections
            draft_sections = self.parse_draft_into_sections(generated_text)

            return draft_sections

        except Exception as e:
            logger.error(f"Exception in generate_draft_sections_with_gpt: {traceback.format_exc()}")
            return None


    def update_ai_suggested_content_count(self, session_id):
        try:
            session_object_id = ObjectId(session_id)
            update_result = self.get_mongo_client_db().find_one_and_update(
                {'user_id': self.user_id, '_id': session_object_id},
                {'$inc': {'ai_suggested_update_count': 1}},
                return_document=ReturnDocument.AFTER
            )

            # Now you can access the final count from the updated document
            final_count = (update_result or {}).get('ai_suggested_update_count', 0)
            return final_count
        except Exception as e:
            logger.error(f"Exception in update_ai_suggested_content_count: {traceback.format_exc()}")
            return 0


    def get_ai_suggested_content_count(self, session_id):
        try:
            session_object_id = ObjectId(session_id)
            draft = self.get_mongo_client_db().find_one(
                {'user_id': self.user_id, '_id': session_object_id},
                {'ai_suggested_update_count': 1},
            )
            return int((draft or {}).get('ai_suggested_update_count', 0))
        except Exception:
            logger.error(f"Exception in get_ai_suggested_content_count: {traceback.format_exc()}")
            return 0


    def fetch_draft_for(self,session_id):
        try:
            # Assuming session_id is unique and corresponds to a single document
            session_object_id = ObjectId(session_id)
            draft = self.get_mongo_client_db().find_one({'user_id':self.user_id, '_id': session_object_id})
            # draft = drafts_collection.find_one({'user_id': session_id})
            if not draft:
                return {'error': 'Draft not found for the provided session_id.'}

            draft_for = draft.get('draft_for', {})
            return {'draft_for': draft_for}
        except Exception as e:
            logger.error(f"Exception in generate_draft_sections_with_gpt: {traceback.format_exc()}")
            return {'error': 'An error occurred while fetching draft details.'}
